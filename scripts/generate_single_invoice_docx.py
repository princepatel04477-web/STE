import sys
import json
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def number_to_words_inr(amount):
    rounded = int(round(amount))
    if rounded == 0:
        return "Zero Rupees Only"

    a = [
        "", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
        "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
        "Seventeen", "Eighteen", "Nineteen"
    ]
    b = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

    def in_words(num):
        if num == 0: return ""
        if num < 20: return a[num] + " "
        if num < 100: return b[num // 10] + (" " + a[num % 10] if num % 10 != 0 else "") + " "
        if num < 1000: return a[num // 100] + " Hundred " + in_words(num % 100)
        if num < 100000: return in_words(num // 1000) + "Thousand " + in_words(num % 1000)
        if num < 10000000: return in_words(num // 100000) + "Lakh " + in_words(num % 100000)
        return in_words(num // 10000000) + "Crore " + in_words(num % 10000000)

    res = in_words(rounded).strip()
    return f"Rupees {res} Only"

def generate_invoice_docx(data, output_file):
    doc = docx.Document()
    
    # Tight page margins for 1-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 1. HEADER
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cell = header_table.cell(0, 0)
    set_cell_background(h_cell, "111827")
    set_cell_margins(h_cell, top=140, bottom=140, left=180, right=180)

    p_h = h_cell.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = p_h.add_run("SURAT TEXTILE EXHIBITION (STE) 2026\n")
    r1.font.name = "Georgia"
    r1.font.size = Pt(15)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(245, 158, 11) # Gold

    r_title = p_h.add_run("TAX INVOICE / ESTIMATE BILL — EXTRA REQUIREMENTS\n")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(255, 255, 255)

    p_addr = h_cell.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_a = p_addr.add_run(
        "Surat textile exhibition\n"
        "5TH FLOOR SHOP NO.-B/503 TEXTILE MARKET MAGOB Surat\n"
        "Phone no.: 9950787787 | Email: surattextileexhibition@gmail.com\n"
        "GSTIN: 24AFOFS4061C1Z3 | State: 24-Gujarat"
    )
    r_a.font.name = "Arial"
    r_a.font.size = Pt(8.5)
    r_a.font.color.rgb = RGBColor(209, 213, 219)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 2. INVOICE META & BILLED TO TABLE
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in meta_table.rows:
        for c in r.cells:
            set_cell_background(c, "F9FAFB")
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)

    brand_name = data.get("brand_name", "Registered Exhibitor")
    mobile = data.get("mobile", "—")
    stall_sqft = data.get("stall_sqft", "200 sq ft")
    fascia_names = data.get("fascia_names", [])
    fascia_str = " | ".join([f for f in fascia_names if f]) if isinstance(fascia_names, list) and any(fascia_names) else ""
    invoice_no = data.get("invoice_no", f"STE/INV/2026/{mobile[-4:] if len(mobile)>=4 else '0001'}")
    date_str = data.get("date", "22-Aug-2026")
    days = int(data.get("days", 2))

    p_bill = meta_table.cell(0, 0).paragraphs[0]
    p_bill.add_run("BILLED TO (EXHIBITOR):\n").font.bold = True
    p_bill.add_run(f"Brand: {brand_name}\nMobile: {mobile}\nStall Size: {stall_sqft}" + (f"\nFascia: {fascia_str}" if fascia_str else "")).font.size = Pt(8.5)

    p_inv = meta_table.cell(0, 1).paragraphs[0]
    p_inv.add_run("INVOICE DETAILS:\n").font.bold = True
    p_inv.add_run(f"Invoice No: {invoice_no}\nDate: {date_str}\nRental Duration: {days} Days\nPlace of Supply: 24-Gujarat").font.size = Pt(8.5)

    meta_table.cell(1, 0).paragraphs[0].add_run("Exhibition: Surat Textile Exhibition 2026").font.size = Pt(8)
    meta_table.cell(1, 1).paragraphs[0].add_run("Venue: SIECC Sarsana, Surat (Sept 12-13, 2026)").font.size = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. ITEMS TABLE
    items = data.get("items", [])
    table = doc.add_table(rows=len(items) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["S.No.", "Item Code & Description", "Rate/Day (INR)", "Days", "Qty", "Amount (INR)"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F2937")
        set_cell_margins(hdr_cells[i], top=70, bottom=70, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(255, 255, 255)

    subtotal = 0
    for idx, item in enumerate(items):
        row_cells = table.rows[idx + 1].cells
        bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
        for col_idx in range(6):
            set_cell_background(row_cells[col_idx], bg)
            set_cell_margins(row_cells[col_idx], top=60, bottom=60, left=80, right=80)

        qty = int(item.get("quantity", item.get("qty", 1)))
        rate = float(item.get("rateInr", item.get("rate_inr", item.get("rate", 0))))
        item_days = int(item.get("days", data.get("days", 2)))
        code = item.get("code", "")
        name = item.get("name", "Extra Item")
        spec = item.get("spec", "")
        
        line_total = rate * qty * item_days
        subtotal += line_total

        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)

        p_desc = row_cells[1].paragraphs[0]
        if code:
            r_c = p_desc.add_run(f"[{code}] ")
            r_c.font.bold = True
            r_c.font.color.rgb = RGBColor(180, 83, 9)
        r_n = p_desc.add_run(name)
        r_n.font.bold = True
        if spec:
            r_s = p_desc.add_run(f" ({spec})")
            r_s.font.color.rgb = RGBColor(107, 114, 128)
        p_desc.runs[0].font.size = Pt(8)

        row_cells[2].text = f"₹{rate:,.0f}"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(8)

        row_cells[3].text = str(item_days)
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].runs[0].font.size = Pt(8)

        row_cells[4].text = str(qty)
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].runs[0].font.size = Pt(8)

        row_cells[5].text = f"₹{line_total:,.0f}"
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[5].paragraphs[0].runs[0].font.bold = True
        row_cells[5].paragraphs[0].runs[0].font.size = Pt(8)

    # 4. CALCULATION TOTALS
    cgst = round(subtotal * 0.09)
    sgst = round(subtotal * 0.09)
    total_gst = cgst + sgst
    grand_total = subtotal + total_gst
    amount_in_words = number_to_words_inr(grand_total)

    tot_table = doc.add_table(rows=4, cols=2)
    tot_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    for r in tot_table.rows:
        for c in r.cells:
            set_cell_margins(c, top=50, bottom=50, left=100, right=100)

    tot_rows = [
        ("Subtotal (Taxable Value):", f"₹{subtotal:,.0f}"),
        ("CGST @ 9% (Gujarat Intra-State):", f"₹{cgst:,.0f}"),
        ("SGST @ 9% (Gujarat Intra-State):", f"₹{sgst:,.0f}"),
        ("Grand Total (Incl. 18% GST):", f"₹{grand_total:,.0f}"),
    ]
    for r_i, (lbl, val) in enumerate(tot_rows):
        c0 = tot_table.cell(r_i, 0)
        c1 = tot_table.cell(r_i, 1)
        c0.text = lbl
        c1.text = val
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if r_i == 3:
            set_cell_background(c0, "FEF3C7")
            set_cell_background(c1, "FEF3C7")
            c0.paragraphs[0].runs[0].font.bold = True
            c0.paragraphs[0].runs[0].font.size = Pt(9.5)
            c1.paragraphs[0].runs[0].font.bold = True
            c1.paragraphs[0].runs[0].font.size = Pt(10)
            c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(180, 83, 9)
        else:
            c0.paragraphs[0].runs[0].font.size = Pt(8.5)
            c1.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 5. AMOUNT IN WORDS
    p_words = doc.add_paragraph()
    p_words.paragraph_format.space_after = Pt(4)
    r_w1 = p_words.add_run("Amount in Words: ")
    r_w1.font.bold = True
    r_w1.font.size = Pt(8.5)
    r_w2 = p_words.add_run(amount_in_words)
    r_w2.font.bold = True
    r_w2.font.size = Pt(8.5)
    r_w2.font.color.rgb = RGBColor(180, 83, 9)

    # 6. BANK ACCOUNT DETAILS & UPI QR CODE
    bank_table = doc.add_table(rows=1, cols=2)
    bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_cell = bank_table.cell(0, 0)
    qr_cell = bank_table.cell(0, 1)
    
    set_cell_background(b_cell, "F3F4F6")
    set_cell_margins(b_cell, top=80, bottom=80, left=120, right=120)
    set_cell_background(qr_cell, "FFFFFF")
    set_cell_margins(qr_cell, top=60, bottom=60, left=80, right=80)

    p_bh = b_cell.paragraphs[0]
    r_bh = p_bh.add_run("COMPANY BANK ACCOUNT DETAILS (NEFT / RTGS / IMPS):\n")
    r_bh.font.bold = True
    r_bh.font.size = Pt(8)
    r_bh.font.color.rgb = RGBColor(17, 24, 39)

    p_bi = b_cell.add_paragraph()
    r_bi = p_bi.add_run(
        "A/C NAME:   SURAT TEXTILE EXHIBITION\n"
        "A/C NO:       183805503938\n"
        "IFSC CODE:  ICIC0001838\n"
        "BANK:         ICICI Bank (Surat Branch)"
    )
    r_bi.font.name = "Courier New"
    r_bi.font.bold = True
    r_bi.font.size = Pt(8)

    # Embed QR Code if present
    qr_paths = [
        os.path.join(os.path.dirname(os.path.dirname(__file__)), "public", "upi_qr.png"),
        os.path.join(os.getcwd(), "public", "upi_qr.png"),
    ]
    qr_file = next((p for p in qr_paths if os.path.exists(p)), None)
    if qr_file:
        p_qr = qr_cell.paragraphs[0]
        p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_qr = p_qr.add_run()
        run_qr.add_picture(qr_file, width=Inches(0.95))
        p_lbl = qr_cell.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p_lbl.add_run("Scan to Pay via UPI\n(GPay / PhonePe / Paytm)")
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(6.5)
        r_lbl.font.color.rgb = RGBColor(180, 83, 9)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 7. SIGNATORY & TERMS
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in sig_table.rows[0].cells:
        set_cell_margins(c, top=80, bottom=80, left=80, right=80)

    p_terms = sig_table.cell(0, 0).paragraphs[0]
    p_terms.add_run("Terms & Conditions:\n").font.bold = True
    p_terms.add_run("• 100% advance payment required.\n• Rent items are for exhibition period.\n• Subject to Surat jurisdiction.").font.size = Pt(7.5)

    p_s = sig_table.cell(0, 1).paragraphs[0]
    p_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_s.add_run("For Surat textile exhibition\n\n\n___________________________\nAuthorized Signatory").font.size = Pt(8)

    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    print(f"Generated invoice at {output_file}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        payload_file = sys.argv[1]
        out_path = sys.argv[2]
        with open(payload_file, "r", encoding="utf-8") as f:
            d = json.load(f)
        generate_invoice_docx(d, out_path)
    else:
        print("Usage: python generate_single_invoice_docx.py <payload.json> <output.docx>")
