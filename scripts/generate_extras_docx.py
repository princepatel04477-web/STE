import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_extras_document(output_path):
    doc = docx.Document()
    
    # Set page margins (0.5 inch / 36 pt for clean layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 1. HEADER TABLE (Branded Header with Company & Event Details)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    set_cell_background(header_cell, "111827") # Dark Midnight Theme
    set_cell_margins(header_cell, top=180, bottom=180, left=200, right=200)

    p1 = header_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("SURAT TEXTILE EXHIBITION (STE) 2026\n")
    run1.font.name = "Georgia"
    run1.font.size = Pt(16)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(245, 158, 11) # Amber / Gold

    run_sub = p1.add_run("OFFICIAL EXTRA AMENITIES & RENTAL RATE CARD\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(255, 255, 255)

    p2 = header_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_addr = p2.add_run(
        "Surat textile exhibition\n"
        "5TH FLOOR SHOP NO.-B/503 TEXTILE MARKET MAGOB Surat\n"
        "Phone no.: 9950787787 | Email: surattextileexhibition@gmail.com\n"
        "GSTIN: 24AFOFS4061C1Z3 | State: 24-Gujarat\n"
        "Exhibition Dates: September 12-13, 2026 | Venue: SIECC Sarsana, Surat"
    )
    r_addr.font.name = "Arial"
    r_addr.font.size = Pt(9)
    r_addr.font.color.rgb = RGBColor(209, 213, 219)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. NOTICE & GUIDELINES
    p_notice = doc.add_paragraph()
    p_notice.paragraph_format.space_after = Pt(8)
    r_n1 = p_notice.add_run("IMPORTANT BOOKING NOTICE:\n")
    r_n1.font.bold = True
    r_n1.font.size = Pt(10)
    r_n1.font.color.rgb = RGBColor(180, 83, 9)

    r_n2 = p_notice.add_run(
        "1. Note: The reference product images are for booking purpose only. The original physical product may slightly vary.\n"
        "2. All rates listed below are EXCLUSIVE OF GST. Applicable 18% GST (CGST 9% + SGST 9% / IGST 18%) will be added to the invoice.\n"
        "3. Strict Order Deadline: 5th September 2026 at 12:00 PM. No modifications permitted after the cutoff."
    )
    r_n2.font.size = Pt(9)

    # 3. EXTRA ITEMS CATALOG TABLE
    items_data = [
        ("DP 01", "Desk Table", "1m × 0.5m × 0.75m", "₹600", "$7", "Per day", "Furniture"),
        ("DP 03", "Glass Round Table", "1m dia × 0.75m", "₹1,400", "$15", "Per day", "Furniture"),
        ("DP 04", "White Chair", "Standard exhibition seating", "₹700", "$8", "Per day", "Furniture"),
        ("DP 05", "Cushioned Chair", "Comfortable cushioned meeting chair", "₹700", "$8", "Per day", "Furniture"),
        ("DP 08", "Sofa Single Seat", "Plush leather armchair", "₹3,000", "$33", "Per day", "Furniture"),
        ("DP 09", "Sofa Double Seat", "2-seater luxury lounge sofa", "₹5,000", "$55", "Per day", "Furniture"),
        ("DP 10", "Sofa Three Seat", "3-seater spacious lounge sofa", "₹6,000", "$66", "Per day", "Furniture"),
        ("DP 11", "Glass Centre Table", "Glass top lounge center table", "₹1,200", "$13", "Per day", "Furniture"),
        ("DP 15", "Brochure Rack", "Acrylic / metal catalogue stand", "₹1,500", "$16", "Per day", "Display & AV"),
        ("DP 20", "Pedestrian Fan", "High power pedestal fan", "₹1,500", "$16", "Per day", "Electrical"),
        ("DP 21", "Glass Shelf", "1m × 0.25m wall shelf", "₹600", "$6", "Per day", "Display & AV"),
        ("DP 22", "Wooden Shelf", "1m × 0.25m wooden shelf", "₹500", "$5", "Per day", "Display & AV"),
        ("DP 24", "Plug Point", "5 / 15 amp single phase socket", "₹250", "$3", "Per day", "Electrical"),
        ("DP 25", "Metal Halide Light", "Spotlight / 50W LED fixture", "₹1,500", "$13", "Per day", "Electrical"),
        ("DP 26", "32\" Plasma Screen with Stand", "32-inch display with floor stand", "₹3,500", "$27", "Per day", "Display & AV"),
        ("DP 27", "Garment Stand", "Display rack for fabric/garment hangers", "₹900", "$7", "Per day", "Display & AV"),
    ]

    table = doc.add_table(rows=len(items_data) + 1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Item Code", "Description", "Specification", "Rate (INR)", "Rate (USD)", "Basis", "Required Qty"]
    col_widths = [Inches(0.9), Inches(1.8), Inches(1.8), Inches(1.0), Inches(0.8), Inches(0.8), Inches(1.0)]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F2937") # Slate 800
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, data in enumerate(items_data):
        row_cells = table.rows[row_idx + 1].cells
        # Alternating background colors
        bg = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx in range(7):
            set_cell_background(row_cells[col_idx], bg)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)

        row_cells[0].text = data[0] # Code
        row_cells[1].text = data[1] # Name
        row_cells[2].text = data[2] # Spec
        row_cells[3].text = data[3] # INR Rate
        row_cells[4].text = data[4] # USD Rate
        row_cells[5].text = data[5] # Basis
        row_cells[6].text = "[   ] Qty" # Requisition placeholder

        # Styling
        p_code = row_cells[0].paragraphs[0]
        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_code.runs[0].font.bold = True
        p_code.runs[0].font.size = Pt(8.5)
        p_code.runs[0].font.color.rgb = RGBColor(180, 83, 9)

        p_name = row_cells[1].paragraphs[0]
        p_name.runs[0].font.bold = True
        p_name.runs[0].font.size = Pt(8.5)

        p_spec = row_cells[2].paragraphs[0]
        p_spec.runs[0].font.size = Pt(8)
        p_spec.runs[0].font.color.rgb = RGBColor(75, 85, 99)

        p_rate = row_cells[3].paragraphs[0]
        p_rate.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_rate.runs[0].font.bold = True
        p_rate.runs[0].font.size = Pt(8.5)
        p_rate.runs[0].font.color.rgb = RGBColor(17, 24, 39)

        p_usd = row_cells[4].paragraphs[0]
        p_usd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_usd.runs[0].font.size = Pt(8)

        p_basis = row_cells[5].paragraphs[0]
        p_basis.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_basis.runs[0].font.size = Pt(8)

        p_qty = row_cells[6].paragraphs[0]
        p_qty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_qty.runs[0].font.size = Pt(8)
        p_qty.runs[0].font.color.rgb = RGBColor(156, 163, 175)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. COMPANY BANK ACCOUNT DETAILS BOX
    bank_table = doc.add_table(rows=1, cols=1)
    bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bank_cell = bank_table.cell(0, 0)
    set_cell_background(bank_cell, "FEF3C7") # Warm Amber Light
    set_cell_margins(bank_cell, top=140, bottom=140, left=180, right=180)

    p_bank_hdr = bank_cell.paragraphs[0]
    r_bkh = p_bank_hdr.add_run("OFFICIAL COMPANY BANK ACCOUNT DETAILS (FOR NEFT / RTGS / IMPS PAYMENT)\n")
    r_bkh.font.name = "Arial"
    r_bkh.font.bold = True
    r_bkh.font.size = Pt(10)
    r_bkh.font.color.rgb = RGBColor(146, 64, 14)

    p_bank_info = bank_cell.add_paragraph()
    r_binfo = p_bank_info.add_run(
        "A/C NAME:      SURAT TEXTILE EXHIBITION\n"
        "A/C NO:          183805503938\n"
        "IFSC CODE:     ICIC0001838\n"
        "BANK NAME:     ICICI Bank (Surat Branch)\n"
        "PAYMENT NOTE:  Please share transaction UTR / payment screenshot on WhatsApp: 9950787787"
    )
    r_binfo.font.name = "Courier New"
    r_binfo.font.bold = True
    r_binfo.font.size = Pt(9.5)
    r_binfo.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 5. EXHIBITOR REQUISITION FORM SECTION
    p_ex_title = doc.add_paragraph()
    r_ext = p_ex_title.add_run("EXHIBITOR BOOKING DETAILS / ORDER FORM:")
    r_ext.font.bold = True
    r_ext.font.size = Pt(10.5)
    r_ext.font.color.rgb = RGBColor(17, 24, 39)

    order_table = doc.add_table(rows=3, cols=2)
    order_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in order_table.rows:
        for c in r.cells:
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    order_table.cell(0, 0).paragraphs[0].add_run("Exhibitor Brand Name: _______________________").font.size = Pt(9)
    order_table.cell(0, 1).paragraphs[0].add_run("Registered Mobile No: _______________________").font.size = Pt(9)
    order_table.cell(1, 0).paragraphs[0].add_run("Allocated Stall Size (Sq Ft): _________________").font.size = Pt(9)
    order_table.cell(1, 1).paragraphs[0].add_run("Contact Person Name: _______________________").font.size = Pt(9)
    order_table.cell(2, 0).paragraphs[0].add_run("Owner Entry Badges Required: ______________").font.size = Pt(9)
    order_table.cell(2, 1).paragraphs[0].add_run("Staff Badges Required: ______________________").font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 5.1 BADGE HOLDER NAMES TABLE
    p_b_title = doc.add_paragraph()
    r_bt = p_b_title.add_run("EXHIBITOR BADGE HOLDER NAMES (TO BE PRINTED ON PASSES):")
    r_bt.font.bold = True
    r_bt.font.size = Pt(9.5)
    r_bt.font.color.rgb = RGBColor(17, 24, 39)

    badge_table = doc.add_table(rows=6, cols=4)
    badge_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_headers = ["S.No.", "Full Name of Person", "Designation / Role", "Mobile No."]
    for c_i, h in enumerate(b_headers):
        cell = badge_table.cell(0, c_i)
        set_cell_background(cell, "F3F4F6")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        if c_i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_i in range(1, 6):
        for c_i in range(4):
            cell = badge_table.cell(r_i, c_i)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            if c_i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(r_i))
                r.font.size = Pt(8)
            elif c_i == 2:
                r = p.add_run("Owner / Sales / Support" if r_i == 1 else "")
                r.font.size = Pt(7.5)
                r.font.color.rgb = RGBColor(156, 163, 175)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 6. SIGNATURE & STAMP
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in sig_table.rows[0].cells:
        set_cell_margins(c, top=100, bottom=100, left=100, right=100)

    p_sig1 = sig_table.cell(0, 0).paragraphs[0]
    p_sig1.add_run("Exhibitor Authorized Signature & Stamp:\n\n\n_____________________________________\nDate: _______________________________").font.size = Pt(8.5)

    p_sig2 = sig_table.cell(0, 1).paragraphs[0]
    p_sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s2 = p_sig2.add_run("For Surat textile exhibition\n\n\n_____________________________________\nAuthorized Signatory (Organizers)").font.size = Pt(8.5)

    # Ensure target directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated Word docx at: {output_path}")

if __name__ == "__main__":
    out1 = os.path.join(os.getcwd(), "public", "downloads", "STE_Extra_Items_Rate_Card_and_Order_Form.docx")
    out2 = os.path.join(os.getcwd(), "public", "STE_EXTRAS_DOCUMENT.docx")
    create_extras_document(out1)
    create_extras_document(out2)
